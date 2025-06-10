#!/usr/bin/env python3
"""
SarahMemorySystemLearn.py
<Version #7.1.2 Enhanced> <Author: Brian Lee Baros> rev. 060920252100

Purpose:
  Ingest and pretrain indexed information for improved recall/responses.
  Supports multiple file types (e.g. .docx, .pdf, .txt, images) and extracts metadata.
  
Notes:
  - Uses Mammoth and Antiword to extract .doc file texts.
  - Populates various databases for later retrieval by the generative AI system.
"""
import sys
import os
import sqlite3
import datetime
import hashlib
import numpy as np
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import spacy
from langdetect import detect
from docx import Document as DocxDocument
import PyPDF2
import subprocess
import threading
import importlib
import SarahMemoryAdvCU
importlib.reload(SarahMemoryAdvCU)
import SarahMemoryGlobals as config
from SarahMemoryGlobals import DATASETS_DIR, TOOLS_DIR #content = extract_text(...)
from SarahMemoryAdvCU import classify_intent
from SarahMemoryAdaptive import advanced_emotional_learning
from SarahMemoryWebSYM import WebSemanticSynthesizer
from SarahMemoryFacialRecognition import add_vector

config.LEARNING_PHASE_ACTIVE = True  # Only TRUE during indexing/learning
# Initialize FAISS and SentenceTransformer
from SarahMemoryGlobals import MODEL_CONFIG, MULTI_MODEL

def get_active_sentence_model():
    from sentence_transformers import SentenceTransformer
    if MULTI_MODEL:
        for model_name, enabled in MODEL_CONFIG.items():
            if enabled:
                try:
                    return SentenceTransformer(model_name)
                except Exception as e:
                    print(f"[MODEL LOAD ERROR] {model_name} failed: {e}")
    return SentenceTransformer("all-MiniLM-L6-v2")  # Fallback default

try:
    import torch
    import faiss
    model = get_active_sentence_model()
    model.to("cuda" if torch.cuda.is_available() else "cpu")
    vector_index = faiss.IndexFlatL2(384)
except Exception as e:
    model = None
    vector_index = None
    print(f"⚠️ Vector engine not initialized: {e}")

# GUI CONFIGURATION
FILETYPE_GROUPS = {
    "doc/docx": ['.doc', '.docx'],
    "pdf": ['.pdf'],
    "text": ['.txt', '.md', '.csv', '.json', '.xml', '.yaml', '.yml'],
    "images": ['.jpg', '.jpeg', '.png', '.bmp', '.tif', '.tiff'],
    "audio": ['.mp3', '.wav', '.ogg', '.flac', '.m4a'],
    "video": ['.mp4', '.mkv', '.mov', '.avi'],
    "code": ['.py', '.ipynb', '.js', '.html', '.css', '.php', '.asp', '.bat', '.sh', '.sql'],
    "odt": ['.odt']
}
def memory_autocorrect():
    """
    Scans system logs and learning cache for flagged or low-confidence entries
    and fixes them by updating the override intent cache for future accuracy.
    """
    import json
    
    override_path = os.path.join(DATASETS_DIR, "intent_override_cache.json")
    cache = {}
    if os.path.exists(override_path):
        try:
            with open(override_path, "r") as f:
                cache = json.load(f)
        except:
            cache = {}

    qa_path = os.path.join(DATASETS_DIR, "functions.db")
    import sqlite3
    conn = sqlite3.connect(qa_path)
    cursor = conn.cursor()
    cursor.execute("SELECT id, query, ai_answer, hit_score FROM qa_cache")
    updates = 0

    for row in cursor.fetchall():
        qa_id, query, answer, hit_score = row
        if hit_score <= 0:
            new_intent = classify_intent(query)
            if query not in cache:
                cache[query] = new_intent
                updates += 1

    with open(override_path, "w") as f:
        json.dump(cache, f, indent=4)

    conn.close()
    print(f"🧠 Autocorrected and cached {updates} misclassified intents.")
    return updates


# ---------------------------------------------------------------------------
# Process files
# ---------------------------------------------------------------------------
def process_files(files, use_mammoth=True, use_antiword=True, learn_registry=False):
    log(f"📁 Processing {len(files)} files from index...")
    total = 0
    for file_path, file_type in files:
        if not os.path.exists(file_path):
            continue
        # Call extract_text with proper flags
        content = extract_text(file_path, use_mammoth=use_mammoth, use_antiword=use_antiword)
        if not content and not file_path.lower().endswith(('.jpg', '.jpeg', '.png')):
            continue
        if content and not is_clean_content(content):
            log(f"🧹 Skipped low-quality or non-English content: {file_path}")
            continue
        category = categorize_text_by_path_or_content(file_path, content or "")
        if not entry_already_learned(category, content or "", file_path):
            if insert_to_database(category, content or "", filepath=file_path):
                total += 1
                log(f"✅ {file_path} => {category}")
    log(f"✅ Completed processing {total} files.")

# ---------------------------------------------------------------------------
# GUI
# ---------------------------------------------------------------------------
PARSERS = []  # Disabled due to Mammoth/Antiword deprecation
# GUI window
class LearnGUI(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("SarahMemory Learning Module")
        self.geometry("800x650")
        
        drives = sorted({os.path.splitdrive(fp)[0] for fp, _ in read_index_db()[0] if os.path.splitdrive(fp)[0]})
        self.selected_drive = tk.StringVar(value="ALL")
        drive_frame = tk.Frame(self)
        drive_frame.pack(pady=5)
        ttk.Label(drive_frame, text="Select Drive to Learn From:").pack(side="left")
        self.drive_menu = ttk.Combobox(drive_frame, textvariable=self.selected_drive, values=["ALL"] + drives)
        self.drive_menu.pack(side="left")

        
        
        self.filetypes = { ext: tk.BooleanVar() for ext in [
            '.doc', '.docx', '.pdf', '.txt', '.jpg', '.png', '.mp3', '.wav',
            '.mp4', '.md', '.csv', '.json', '.xml', '.yaml', '.yml', '.odt',
            '.jpeg', '.bmp', '.tif', '.tiff', '.ogg', '.flac', '.m4a',
            '.mkv', '.mov', '.avi', '.py', '.ipynb', '.js', '.html', '.css',
            '.php', '.asp', '.bat', '.sh', '.sql'
        ]}
        self.mammoth_enabled = tk.BooleanVar(value=True)
        self.antiword_enabled = tk.BooleanVar(value=True)
        self.registry_enabled = tk.BooleanVar(value=False)
        self.init_gui()

    def init_gui(self):
        frame = tk.Frame(self)
        frame.pack(pady=10)

        ttk.Label(frame, text="Select File Types to Learn From:").pack()
        types_frame = tk.Frame(frame)
        types_frame.pack()

        for idx, (ext, var) in enumerate(self.filetypes.items()):
            cb = ttk.Checkbutton(types_frame, text=ext, variable=var)
            cb.grid(row=idx//6, column=idx % 6, sticky='w')

        options = tk.Frame(frame)
        options.pack(pady=10)
        ttk.Checkbutton(options, text="Use Mammoth", variable=self.mammoth_enabled).pack(side="left")
        ttk.Checkbutton(options, text="Use Antiword", variable=self.antiword_enabled).pack(side="left")
        ttk.Checkbutton(options, text="Learn Registry", variable=self.registry_enabled).pack(side="left")

        action_frame = tk.Frame(frame)
        action_frame.pack(pady=10)
        ttk.Button(action_frame, text="Learn From Indexed Files", command=self.start_indexed_learning).pack(side="left", padx=10)
        ttk.Button(action_frame, text="Browse and Learn Single File", command=self.learn_single_file).pack(side="left", padx=10)
        ttk.Button(action_frame, text="Import Dataset (.db/.csv/.xml)", command=self.import_external_dataset).pack(side="left", padx=10)

        self.console = scrolledtext.ScrolledText(self, wrap=tk.WORD, height=20)
        self.console.pack(fill="both", expand=True)

    def import_external_dataset(self):
        filepath = filedialog.askopenfilename(filetypes=[("Data Files", "*.db *.csv *.xml *.json *.xlsx")])
        if not filepath:
            return
        try:
            if filepath.endswith(".db"):
                self.print_console(f"🧠 Importing database: {filepath}")
                self.absorb_sqlite(filepath)
            elif filepath.endswith(".csv"):
                self.print_console(f"🧠 Importing CSV: {filepath}")
                self.absorb_csv(filepath)
            elif filepath.endswith(".json"):
                self.print_console(f"🧠 Importing JSON: {filepath}")
                self.absorb_json(filepath)
            elif filepath.endswith(".xlsx"):
                self.print_console(f"🧠 Importing Excel: {filepath}")
                self.absorb_excel(filepath)
            elif filepath.endswith(".xml"):
                self.print_console(f"🧠 Importing XML: {filepath}")
                self.absorb_xml(filepath)
        except Exception as e:
            self.print_console(f"❌ Import failed: {e}")

    def print_console(self, msg):
        self.console.insert(tk.END, msg + "\n")
        self.console.see(tk.END)

    def start_indexed_learning(self):
        from SarahMemorySystemLearn import read_index_db
        files, registry = read_index_db()
        selected_exts = [ext for ext, var in self.filetypes.items() if var.get()]
        drive = self.selected_drive.get()
        if drive != "ALL":
            files = [(f, t) for f, t in files if f.upper().startswith(drive.upper())]
        filtered = [(f, t) for f, t in files if os.path.splitext(f)[1].lower() in selected_exts]
        self.print_console(f"Processing {len(filtered)} files from drive {drive}...")
        process_files(filtered, use_mammoth=self.mammoth_enabled.get(), 
                      use_antiword=self.antiword_enabled.get(),
                      learn_registry=self.registry_enabled.get())
        self.print_console("✅ Finished Learning From Indexed Files.")

    def process_indexed_files(self, exts):
        from SarahMemorySystemLearn import read_index_db, process_files
        files, registry = read_index_db()
        filtered = [(path, typ) for path, typ in files if os.path.splitext(path)[1].lower() in exts]
        self.print_console(f"Processing {len(filtered)} files...")
        process_files(filtered, use_mammoth=self.mammoth_enabled.get(), use_antiword=self.antiword_enabled.get(), learn_registry=self.registry_enabled.get())
        self.print_console("✅ Finished Learning From Indexed Files.")

    def learn_single_file(self):
        path = filedialog.askopenfilename()
        if not path:
            return
        from SarahMemorySystemLearn import extract_text, categorize_text_by_path_or_content, entry_already_learned, insert_to_database
        ext = os.path.splitext(path)[1].lower()
        content = extract_text(path, use_mammoth=self.mammoth_enabled.get(), use_antiword=self.antiword_enabled.get())
        if not content:
            self.print_console(f"❌ Unable to extract from: {path}")
            return
        category = categorize_text_by_path_or_content(path, content)
        if not entry_already_learned(category, content, path):
            insert_to_database(category, content, path)
            self.print_console(f"✅ Learned: {path} => {category}")
        else:
            self.print_console(f"⏭️ Skipped already-learned file: {path}")

log_console = None
def launch_learning_gui():
    root = tk.Tk()
    root.title("SarahMemory System Learn Engine By Brian Lee Baros")
    root.geometry("900x750")

    selected_drive = tk.StringVar()
    # Add default values for mammoth and antiword
    parser_vars = {"mammoth": tk.BooleanVar(value=True), "antiword": tk.BooleanVar(value=True)}
    registry_var = tk.IntVar(value=0)
    group_vars = {k: tk.IntVar(value=1) for k in FILETYPE_GROUPS}


    def absorb_sqlite(self, filepath):
        try:
            ext_conn = sqlite3.connect(filepath)
            ext_cursor = ext_conn.cursor()
            ext_cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
            tables = ext_cursor.fetchall()
            for (table,) in tables:
                try:
                    ext_cursor.execute(f"SELECT * FROM {table}")
                    rows = ext_cursor.fetchall()
                    col_names = [desc[0] for desc in ext_cursor.description]
                    for row in rows:
                        row_dict = dict(zip(col_names, row))
                        content = " ".join(str(v) for v in row_dict.values())
                        if content.strip():
                            category = categorize_text_by_path_or_content(filepath, content)
                            insert_to_database(category, content, filepath)
                except Exception as table_error:
                    self.print_console(f"⚠️ Skipped table '{table}': {table_error}")
            ext_conn.close()
        except Exception as e:
            self.print_console(f"❌ SQLite import error: {e}")

    def absorb_csv(self, filepath):
        try:
            import csv
            with open(filepath, newline='', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile)
                for row in reader:
                    content = " ".join(str(v) for v in row.values())
                    if content.strip():
                        category = categorize_text_by_path_or_content(filepath, content)
                        insert_to_database(category, content, filepath)
        except Exception as e:
            self.print_console(f"❌ CSV import error: {e}")


    def absorb_json(self, filepath):
        try:
            import json
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                entries = [data]
            elif isinstance(data, list):
                entries = data
            else:
                self.print_console("⚠️ Unrecognized JSON structure.")
                return
            for entry in entries:
                if isinstance(entry, dict):
                    content = " ".join(str(v) for v in entry.values())
                    if content.strip():
                        category = categorize_text_by_path_or_content(filepath, content)
                        insert_to_database(category, content, filepath)
        except Exception as e:
            self.print_console(f"❌ JSON import error: {e}")

    def absorb_excel(self, filepath):
        try:
            import pandas as pd
            df = pd.read_excel(filepath)
            for _, row in df.iterrows():
                content = " ".join(str(v) for v in row.values)
                if content.strip():
                    category = categorize_text_by_path_or_content(filepath, content)
                    insert_to_database(category, content, filepath)
        except Exception as e:
            self.print_console(f"❌ Excel import error: {e}")
    def absorb_xml(self, filepath):
        try:
            import xml.etree.ElementTree as ET
            tree = ET.parse(filepath)
            root = tree.getroot()
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    content = elem.text.strip()
                    category = categorize_text_by_path_or_content(filepath, content)
                    insert_to_database(category, content, filepath)
        except Exception as e:
            self.print_console(f"❌ XML import error: {e}")

    

    def toggle_all():
        state = all(v.get() for v in group_vars.values())
        for v in group_vars.values():
            v.set(0 if state else 1)

    tk.Label(root, text="Select Drive:").pack(pady=5)
    drive_combo = ttk.Combobox(root, textvariable=selected_drive)
    drive_combo['values'] = ["ALL"] + get_available_drives()
    drive_combo.current(0)
    drive_combo.pack(pady=5)

    fbox = tk.LabelFrame(root, text="Select File Types to Learn From")
    fbox.pack(pady=10, fill="x")
    for idx, (label, _) in enumerate(FILETYPE_GROUPS.items()):
        cb = tk.Checkbutton(fbox, text=label, variable=group_vars[label])
        cb.grid(row=idx//4, column=idx%4, sticky='w', padx=10, pady=5)

    ctrl_frame = tk.Frame(root)
    ctrl_frame.pack(pady=10)
    tk.Checkbutton(ctrl_frame, text="Registry", variable=registry_var).grid(row=0, column=0, padx=10)


    def browse_single_file():
        filepath = filedialog.askopenfilename()
        if not filepath:
            return
        file_ext = os.path.splitext(filepath)[1].lower()
        use_mammoth = parser_vars["mammoth"].get()
        use_antiword = parser_vars["antiword"].get()
        registry_flag = registry_var.get()
        process_files([(filepath, file_ext)], use_mammoth, use_antiword, registry_flag)
        messagebox.showinfo("Single File Processed", f"Processed {os.path.basename(filepath)}")

    def start_learning():
        files, registry = read_index_db()
        active_exts = [ext for group, active in group_vars.items() if active.get() for ext in FILETYPE_GROUPS[group]]
        drive = selected_drive.get()
        if drive != "ALL":
            files = [(f, t) for f, t in files if f.upper().startswith(drive.upper())]
        filtered = [(f, t) for f, t in files if os.path.splitext(f)[1].lower() in active_exts]
        if not filtered:
            messagebox.showwarning("No Files", "No matching files found in index.")
            return
        process_files(filtered, 
                      use_mammoth=parser_vars["mammoth"].get(), 
                      use_antiword=parser_vars["antiword"].get(), 
                      learn_registry=registry_var.get())
        messagebox.showinfo("Learning Complete", f"Processed {len(filtered)} files.")

    tk.Button(root, text="Browse & Learn One File", command=lambda: None, bg="blue", fg="white", width=25).pack(pady=5)
    tk.Button(root, text="Start Learning", command=start_learning, bg="green", fg="white", width=25).pack(pady=5)
    tk.Button(root, text="Select/Deselect All File Types", command=toggle_all, bg="gray", fg="white", width=25).pack(pady=5)
    tk.Button(root, text="Exit", command=root.destroy).pack(pady=5)

    root.mainloop()

def launch_learning_gui():
    gui = LearnGUI()
    gui.mainloop()

def get_available_drives():
    drives = set()
    if os.path.exists(INDEX_DB):
        conn = sqlite3.connect(INDEX_DB)
        cur = conn.cursor()
        cur.execute("SELECT file_path FROM file_index")
        for (fp,) in cur.fetchall():
            root = os.path.splitdrive(fp)[0]
            if root:
                drives.add(root)
        conn.close()
    return sorted(drives)

# # Define antiword path globally from tools directory
ANTIWORD_PATH = os.path.join(TOOLS_DIR, "antiword", "antiword.exe")

nlp = spacy.load("en_core_web_sm")
nlp.max_length = 1000000000  # Allows up to 2 million characters
# DB Paths
PERSONALITY_DB = os.path.join(DATASETS_DIR, "personality1.db")
FUNCTIONS_DB = os.path.join(DATASETS_DIR, "functions.db")
SOFTWARE_DB = os.path.join(DATASETS_DIR, "software.db")
PROGRAMMING_DB = os.path.join(DATASETS_DIR, "programming.db")
AVATAR_DB = os.path.join(DATASETS_DIR, "avatar.db")
SYSTEM_LOGS_DB = os.path.join(DATASETS_DIR, "system_logs.db")
INDEX_DB = os.path.join(DATASETS_DIR, "system_index.db")


# ---------------------------------------------------------------------------
# Helper functions, logging, and database operations
# ---------------------------------------------------------------------------
def log(msg):
    timestamp = datetime.datetime.now().isoformat()
    print(f"[{timestamp}] {msg}")
    try:
        conn = sqlite3.connect(os.path.join(DATASETS_DIR, "system_logs.db"))
        conn.execute("""
            CREATE TABLE IF NOT EXISTS logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp TEXT,
                level TEXT,
                source TEXT,
                message TEXT
            )
        """)
        conn.execute("""
            INSERT INTO logs (timestamp, level, source, message)
            VALUES (?, ?, ?, ?)
        """, (timestamp, "INFO", "SystemLearn", msg))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"[LogError] Failed to log to system_logs.db: {e}")
# --- VECTOR INDEX PERSISTENCE CHECK ---
VECTOR_INDEX_PATH = os.path.join(DATASETS_DIR, "vectors.index")
VECTOR_HASH_PATH = os.path.join(DATASETS_DIR, "vector_hash.md5")
FORCE_REBUILD = False

def compute_dataset_hash(directory):
    hash_md5 = hashlib.md5()
    for root, _, files in os.walk(directory):
        for name in sorted(files):
            path = os.path.join(root, name)
            try:
                with open(path, 'rb') as f:
                    while chunk := f.read(8192):
                        hash_md5.update(chunk)
            except Exception:
                continue
    return hash_md5.hexdigest()

# Check vector index and hash
if os.path.exists(VECTOR_INDEX_PATH) and not FORCE_REBUILD:
    current_hash = compute_dataset_hash(DATASETS_DIR)
    if os.path.exists(VECTOR_HASH_PATH):
        with open(VECTOR_HASH_PATH, 'r') as f:
            cached_hash = f.read().strip()
        if current_hash == cached_hash:
            log("✅ Vector index and dataset hash match. Skipping vector rebuild.")
            sys.exit(0)
    # Update hash file
    with open(VECTOR_HASH_PATH, 'w') as f:
        f.write(current_hash)
else:
    log("⚠️ Vector index missing or rebuild forced. Proceeding with vector rebuild.")
# Index reader
# ⬇️ NEW FUNCTION: Extract named entities from document

def extract_named_entities(text):
    doc = nlp(text)
    return [(ent.text, ent.label_) for ent in doc.ents if ent.label_ in ["PERSON", "GPE", "ORG"]]

# ⬇️ NEW FUNCTION: Detect language of content

def detect_language(text):
    try:
        return detect(text)
    except:
        return "unknown"

# ⬇️ NEW FUNCTION: Extract metadata from PDFs and DOCX

def extract_doc_metadata(file_path):
    try:
        if file_path.endswith(".pdf"):
            with open(file_path, "rb") as f:
                pdf = PyPDF2.PdfReader(f)
                meta = pdf.metadata
                return {
                    "created": str(meta.get('/CreationDate', '')),
                    "author": meta.get('/Author', '')
                }
        elif file_path.endswith(".docx"):
            doc = DocxDocument(file_path)
            core_props = doc.core_properties
            return {
                "created": str(core_props.created),
                "author": str(core_props.author)
            }
    except Exception as e:
        print(f"[ExtractDocMetadataError] {file_path}: {e}")
    return {}

def extract_text(file_path, use_mammoth=True, use_antiword=True):
    text = ""
    try:
        if file_path.endswith(".docx"):
            from docx import Document
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif file_path.endswith(".pdf"):
            import PyPDF2
            with open(file_path, "rb") as f:
                pdf = PyPDF2.PdfReader(f)
                for page in pdf.pages:
                    text += page.extract_text() or ""
        elif file_path.endswith(".doc"):
            # Try Mammoth if enabled
            if use_mammoth:
                try:
                    import mammoth
                    with open(file_path, "rb") as doc_file:
                        result = mammoth.convert_to_text(doc_file)
                    text = result.value
                except Exception as e:
                    print(f"[MammothError] {file_path}: {e}")
            # Fallback to Antiword if enabled and Mammoth failed
            if not text and use_antiword:
                import subprocess, os
                env = os.environ.copy()
                env["HOME"] = os.getcwd()
                for mapping_code in ["8859-1", "8859-2", "8859-9"]:
                    try:
                        mapping_path = os.path.join(TOOLS_DIR, "antiword", f"{mapping_code}.txt")
                        antiword_exe = os.path.join(TOOLS_DIR, "antiword", "antiword.exe")
                        result = subprocess.run(
                            [antiword_exe, f"-m{mapping_path}", file_path],
                            stdout=subprocess.PIPE,
                            stderr=subprocess.PIPE,
                            text=True,
                            env=env
                        )
                        if result.returncode == 0:
                            text = result.stdout.strip()
                            break
                        else:
                            print(f"[AntiwordSubprocessError] [{mapping_code}] {file_path}: {result.stderr}")
                    except Exception as e2:
                        print(f"[AntiwordError] [{mapping_code}] {file_path}: {e2}")
        elif file_path.endswith((".txt", ".xml")):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
    except Exception as e:
        print(f"[ExtractTextError] {file_path}: {e}")
    return text.strip()

# ⬇️ INJECTION POINT IN `insert_to_database()` — Add:
   
def read_index_db():
    INDEX_DB = os.path.join(DATASETS_DIR, "system_index.db")
    if not os.path.exists(INDEX_DB):
        log("❌ system_index.db not found!")
        return [], []
    conn = sqlite3.connect(INDEX_DB)
    cur = conn.cursor()
    cur.execute("SELECT file_path, file_type FROM file_index")
    files = cur.fetchall()
    cur.execute("SELECT root_key, key_path, value_name, value_data FROM registry_index")
    registry = cur.fetchall()
    conn.close()
    return files, registry

# ---------------------------------------------------------------------------
# Categorization, vectorization, and insertion functions
# (Keep these functions as in your original code)
# ---------------------------------------------------------------------------
def categorize_text_by_path_or_content(path, text):
    intent = classify_intent(text)
    if WebSemanticSynthesizer.is_math_query(text):
        return "programming"
    p = path.lower()
    if any(kw in p for kw in ["avatar", "profile", ".jpg", ".png"]):
        return "avatar"
    if any(kw in p for kw in ["reminder", "calendar", "schedule", ".ics", ".vcf", ".contact"]):
        return "reminders"
    if any(kw in p for kw in ["userprefs", "settings", "theme", ".ini", ".profile", "ai_name"]):
        return "user_profile"
    if any(p.endswith(ext) for ext in [".bat", ".cmd", ".sh", "notepad.exe", "mspaint.exe", "control.exe"]):
        return "windows10"
    if intent == "command":
        return "functions"
    if intent == "question" or "install" in p or "program files" in p:
        return "software"
    if intent == "statement" and any(x in p for x in ["code", ".py", ".cpp"]):
        return "programming"
    return "personality"
# Vector conversion (basic)

def text_to_vector(text):
    try:
        from sentence_transformers import SentenceTransformer
        from SarahMemoryGlobals import MODEL_CONFIG, MULTI_MODEL

        # Dynamically choose active sentence model
        def get_active_sentence_model():
            if MULTI_MODEL:
                for model_name, enabled in MODEL_CONFIG.items():
                    if enabled:
                        try:
                            return SentenceTransformer(model_name)
                        except Exception as e:
                            print(f"[MODEL LOAD ERROR] {model_name} failed: {e}")
            return SentenceTransformer("all-MiniLM-L6-v2")  # Fallback default

        import torch
        model = get_active_sentence_model()
        device = "cuda" if torch.cuda.is_available() else "cpu"
        model.to(device)

        embedding = model.encode(text, convert_to_tensor=True).cpu().numpy()
        return embedding

    except Exception as e:
        log(f"⚠️ SentenceTransformer fallback engaged due to error: {e}")
        words = text.lower().split()
        vector = np.zeros(128, dtype=np.float32)
        for i, word in enumerate(words[:128]):
            vector[i] = float(len(word)) % 10
        norm = np.linalg.norm(vector)
        return vector / norm if norm > 0 else vector


def extract_exif_info(image_path):
    data = {
        "path": image_path,
        "datetime": None,
        "camera": None,
        "gps_lat": None,
        "gps_lon": None
    }
    try:
        image = Image.open(image_path)
        exif_data = image._getexif()
        if not exif_data:
            return data
        for tag, value in exif_data.items():
            key = TAGS.get(tag, tag)
            if key == 'DateTimeOriginal':
                data["datetime"] = value
            elif key == 'Model':
                data["camera"] = value
            elif key == 'GPSInfo':
                gps_info = {}
                for t in value:
                    sub = GPSTAGS.get(t, t)
                    gps_info[sub] = value[t]
                def convert_gps(gps):
                    d, m, s = gps
                    return d[0]/d[1] + m[0]/m[1]/60 + s[0]/s[1]/3600
                if 'GPSLatitude' in gps_info and 'GPSLatitudeRef' in gps_info:
                    lat = convert_gps(gps_info['GPSLatitude'])
                    if gps_info['GPSLatitudeRef'] != 'N':
                        lat = -lat
                    data['gps_lat'] = lat
                if 'GPSLongitude' in gps_info and 'GPSLongitudeRef' in gps_info:
                    lon = convert_gps(gps_info['GPSLongitude'])
                    if gps_info['GPSLongitudeRef'] != 'E':
                        lon = -lon
                    data['gps_lon'] = lon
    except Exception as e:
        log(f"EXIF extraction failed for {image_path}: {e}")
    return data

# Insert metadata-aware info into DBs

def entry_already_learned(category, content, filepath=""):
    try:
        if category == "avatar":
            return False  # Always store EXIF
        elif category == "functions":
            db = sqlite3.connect(FUNCTIONS_DB)
            cur = db.cursor()
            cur.execute("SELECT id FROM functions WHERE description = ?", (content[:255],))
            exists = cur.fetchone()
            db.close()
            return exists is not None
        elif category == "software":
            db = sqlite3.connect(SOFTWARE_DB)
            cur = db.cursor()
            cur.execute("SELECT app_name FROM software_apps WHERE app_name = ?", ("System Indexed Software",))
            exists = cur.fetchone()
            db.close()
            return exists is not None
        elif category == "programming":
            db = sqlite3.connect(PROGRAMMING_DB)
            cur = db.cursor()
            cur.execute("SELECT id FROM knowledge_base WHERE content = ?", (content,))
            exists = cur.fetchone()
            db.close()
            return exists is not None
        elif category == "reminders":
            db = sqlite3.connect(os.path.join(DATASETS_DIR, "reminders.db"))
            cur = db.cursor()
            cur.execute("SELECT id FROM reminders WHERE description = ?", (content[:255],))
            exists = cur.fetchone()
            db.close()
            return exists is not None
        elif category == "user_profile":
            db = sqlite3.connect(os.path.join(DATASETS_DIR, "user_profile.db"))
            cur = db.cursor()
            cur.execute("SELECT id FROM user_preferences WHERE advanced_metrics = ?", (content[:255],))
            exists = cur.fetchone()
            db.close()
            return exists is not None
        elif category == "windows10":
            db = sqlite3.connect(os.path.join(DATASETS_DIR, "windows10.db"))
            cur = db.cursor()
            cur.execute("SELECT id FROM os_commands WHERE description = ?", (content[:255],))
            exists = cur.fetchone()
            db.close()
            return exists is not None
        elif category == "personality":
            db = sqlite3.connect(PERSONALITY_DB)
            cur = db.cursor()
            cur.execute("SELECT id FROM responses WHERE response = ?", (content[:500],))
            exists = cur.fetchone()
            db.close()
            return exists is not None
        return False
    except Exception as e:
        log(f"[DupCheckError] {e}")
        return False

def insert_to_database(category, content, filepath=""):
    now = datetime.datetime.now().isoformat()
    try:
        if category == "avatar" and filepath.lower().endswith(('.jpg', '.jpeg', '.png')):
            info = extract_exif_info(filepath)
            db = sqlite3.connect(AVATAR_DB)
            db.execute("""
                CREATE TABLE IF NOT EXISTS photo_metadata (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_path TEXT, taken_date TEXT, gps_lat REAL, gps_lon REAL, camera TEXT
                )
            """)
            db.execute("""
                INSERT INTO photo_metadata (file_path, taken_date, gps_lat, gps_lon, camera)
                VALUES (?, ?, ?, ?, ?)
            """, (info['path'], info['datetime'], info['gps_lat'], info['gps_lon'], info['camera']))
            db.commit()
            db.close()
            log(f"🖼️ Photo metadata saved for {filepath}")

        elif category == "functions":
            db = sqlite3.connect(FUNCTIONS_DB)
            db.execute("""
                      INSERT INTO functions (function_name, description, is_enabled, user_input, timestamp)
                      VALUES (?, ?, ?, ?, ?)
                    """, (
                        func_name,                 # Full function name parsed
                        clean_text[:255],          # Cleaned up function description (max 255 chars)
                        True,                      # is_enabled set to True by default
                        func_name,                 # Using function name as default user_input value
                        datetime.now().isoformat() # Current timestamp
                    ))
            db.commit()
            db.close()
        elif category == "software":
            db = sqlite3.connect(SOFTWARE_DB)
            db.execute("""
                INSERT OR IGNORE INTO software_apps (app_name, category, path, is_installed)
                VALUES (?, ?, ?, 1)
            """, ("System Indexed Software", "utility", "system-indexed"))
            db.commit()
            db.close()
        elif category == "programming":
            db = sqlite3.connect(PROGRAMMING_DB)
            db.execute("""
                INSERT INTO knowledge_base (category, content)
                VALUES (?, ?)
            """, ("system-indexed", content))
            db.commit()
            db.close()
        elif category == "reminders":
            db = sqlite3.connect(os.path.join(DATASETS_DIR, "reminders.db"))
            db.execute("""
                CREATE TABLE IF NOT EXISTS reminders (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    title TEXT,
                    description TEXT,
                    datetime TEXT,
                    repeat TEXT DEFAULT 'none',
                    priority INTEGER DEFAULT 0,
                    active BOOLEAN DEFAULT 1
                )
            """)
            db.execute("""
                INSERT INTO reminders (title, description, datetime)
                VALUES (?, ?, ?)
            """, ("Indexed Reminder", content[:255], now))
            db.commit()
            db.close()
        elif category == "user_profile":
            db = sqlite3.connect(os.path.join(DATASETS_DIR, "user_profile.db"))
            db.execute("""
                CREATE TABLE IF NOT EXISTS user_preferences (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ai_name TEXT,
                    voice_pitch REAL,
                    voice_speed REAL,
                    theme TEXT,
                    language TEXT,
                    accessibility_mode BOOLEAN,
                    advanced_metrics TEXT
                )
            """)
            db.execute("""
                INSERT INTO user_preferences (ai_name, voice_pitch, voice_speed, theme, language, accessibility_mode, advanced_metrics)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, ("Sarah", 1.0, 1.0, "dark", "en", 0, content[:255]))
            db.commit()
            db.close()
        elif category == "windows10":
            db = sqlite3.connect(os.path.join(DATASETS_DIR, "windows10.db"))
            db.execute("""
                CREATE TABLE IF NOT EXISTS os_commands (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    command TEXT,
                    description TEXT,
                    version TEXT CHECK(version IN ('10', '11'))
                )
            """)
            db.execute("""
                INSERT INTO os_commands (command, description, version)
                VALUES (?, ?, ?)
            """, ("System Indexed Command", content[:255], "10"))
            db.commit()
            db.close()
        elif category == "personality":
            emo = advanced_emotional_learning(content)
            mood_summary = f"Emotional balance: {emo['emotional_balance']}, Engagement: {emo['engagement']}"
            db = sqlite3.connect(PERSONALITY_DB)
            db.execute("""
                INSERT INTO responses (intent, response, tone, complexity, timestamp)
                VALUES (?, ?, ?, ?, ?)
            """, ("statement", content[:500], "neutral", "student", now))
            db.commit()
            db.close()
            log(f"💬 Emotion captured: {mood_summary}")
        
        if content:
        # Entity tagging
            named_entities = extract_named_entities(content)
            for name, label in named_entities:
                db = sqlite3.connect(PERSONALITY_DB)
                db.execute("""
                    CREATE TABLE IF NOT EXISTS named_references (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        name TEXT, type TEXT, source TEXT, timestamp TEXT
                    )""")
                db.execute("INSERT INTO named_references (name, type, source, timestamp) VALUES (?, ?, ?, ?)",
                        (name, label, filepath, now))
                db.commit()
                db.close()

            # Language tagging
            lang = detect_language(content)
            db = sqlite3.connect(PERSONALITY_DB)
            db.execute("""
                CREATE TABLE IF NOT EXISTS language_references (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    path TEXT, lang TEXT, timestamp TEXT
                )""")
            db.execute("INSERT INTO language_references (path, lang, timestamp) VALUES (?, ?, ?)",
                    (filepath, lang, now))
            db.commit()
            db.close()

        # Metadata tagging
        if filepath.endswith(".pdf") or filepath.endswith(".docx"):
            meta = extract_doc_metadata(filepath)
            db = sqlite3.connect(PERSONALITY_DB)
            db.execute("""
                CREATE TABLE IF NOT EXISTS doc_metadata (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    path TEXT, created TEXT, author TEXT, timestamp TEXT
                )""")
            db.execute("INSERT INTO doc_metadata (path, created, author, timestamp) VALUES (?, ?, ?, ?)",
                       (filepath, meta.get("created", ""), meta.get("author", ""), now))
            db.commit()
            db.close()
        
        if content:
            vector = text_to_vector(content)
            add_vector(vector)
        return True
    except Exception as e:
        log(f"❌ Failed to insert to {category}: {e}")
        return False

def is_clean_content(text):
    if not text or len(text.strip()) < 5:
        return False
    try:
        if detect(text) != 'en':
            return False
    except:
        return False
    doc = nlp(text)
    if not any(t.pos_ in {"NOUN", "VERB", "PROPN"} for t in doc):
        return False
    return True

# # def process_files(files, use_mammoth=True, use_antiword=True, learn_registry=False):
    log(f"📁 Processing {len(files)} files from index...")
    total = 0
    for file_path, file_type in files:
        if not os.path.exists(file_path):
            continue
            content = extract_text(file_path, use_mammoth=use_mammoth, use_antiword=use_antiword)
        if not content and not file_path.lower().endswith(('.jpg', '.jpeg', '.png')):
            continue
        if content and not is_clean_content(content):
            log(f"🧹 Skipped low-quality or non-English content: {file_path}")
            continue
        category = categorize_text_by_path_or_content(file_path, content or "")
        if not entry_already_learned(category, content or "", file_path) and insert_to_database(category, content or "", filepath=file_path):
            total += 1
            log(f"✅ {file_path} => {category}")
    log(f"✅ Completed processing {total} files.")



# ---------------------------------------------------------------------------
# Main Loop
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    launch_learning_gui()

    from SarahMemoryGlobals import IMPORT_OTHER_DATA_LEARN
    if not IMPORT_OTHER_DATA_LEARN:
        log("🛑 SystemLearn vector embedding skipped: IMPORT_OTHER_DATA_LEARN is False.")
        sys.exit(0)

    log("🚀 Starting SarahMemory System Learning Phase...")
    files, registry = read_index_db()
    if files:
        process_files(files)  # Process all indexed files using default extraction flags
    else:
        log("No files found in system index.")
    log("🎓 Knowledge parsed and absorbed into core Class 1 datasets.")

    # NEW: QA Cache Initialization
    try:
        qa_conn = sqlite3.connect(os.path.join(config.DATASETS_DIR, "ai_learning.db"))
        qa_conn.execute("""
            CREATE TABLE IF NOT EXISTS qa_cache (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                query TEXT,
                ai_answer TEXT,
                hit_score INTEGER,
                feedback TEXT,
                timestamp TEXT
            )
        """)
        for file_path, file_type in files:
            content = global_extract_text(file_path)
            if not content:
                continue
            category = categorize_text_by_path_or_content(file_path, content)
            if category in {"functions", "programming", "personality"}:
                qa_conn.execute("""
                    INSERT INTO qa_cache (query, ai_answer, hit_score, feedback, timestamp)
                    VALUES (?, ?, ?, ?, ?)
                """, (content[:200], content[:300], 0, "ungraded", datetime.datetime.now().isoformat()))
        qa_conn.commit()
        qa_conn.close()
        log("📚 QA cache initialized with training examples for future grading.")
    except Exception as e:
        log(f"❌ Failed to initialize qa_cache: {e}")

    # NEW: Process registry software entries...
    software_entries = [(rk, kp, vn, vd) for rk, kp, vn, vd in registry if vn in ('friendly_name', 'exe_path')]
    software_map = {}
    for rk, kp, vn, vd in software_entries:
        key = kp.lower()
        if key not in software_map:
            software_map[key] = {}
        software_map[key][vn] = vd

    if software_map:
        db = sqlite3.connect(os.path.join(DATASETS_DIR, "software.db"))
        db.execute("""
            CREATE TABLE IF NOT EXISTS software_apps (
                app_name TEXT PRIMARY KEY,
                category TEXT,
                path TEXT,
                is_installed BOOLEAN DEFAULT 1
            )
        """)
        for entry in software_map.values():
            name = entry.get("friendly_name")
            path = entry.get("exe_path")
            if name and path:
                db.execute("""
                    INSERT OR REPLACE INTO software_apps (app_name, category, path, is_installed)
                    VALUES (?, ?, ?, 1)
                """, (name, "system", path))
        db.commit()
        db.close()
        log(f"🔍 Mapped {len(software_map)} registry software entries into software.db")

